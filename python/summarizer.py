import sys
import io
import os
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
import pyperclip
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from PyPDF2 import PdfReader

# Redirect stderr to null to suppress warnings
sys.stderr = open(os.devnull, 'w')

# Access your API key as an environment variable.
load_dotenv()
GOOGLE_API_KEY = os.getenv("API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

# Choose a model that's appropriate for your use case.
model = genai.GenerativeModel(
    'gemini-1.5-flash',
    generation_config=genai.GenerationConfig(max_output_tokens=2000, temperature=0.9,)) #TEMPERATURE controls the randomness of the output. Use higher values for more creative responses, and lower values for more deterministic responses. Values can range from [0.0, 2.0]. also MaxOutputTokens sets the maximum number of tokens to include in a candidate.

def get_input():
    choice = input("Do you want to (1) Enter text manually or (2) Upload a file? (Enter 1 or 2): ")
    if choice == '1':
        print("\n" + "=" * 50)
        print("Enter the text you want to summarize below:")
        print("(Type your text and press Enter when finished)")
        print("-" * 50)
        text = input()
        # text = sys.stdin.read()
        print("=" * 50)
        return text
    elif choice == '2':
        filepath = input("Enter the file path: ")
        return read_file(filepath)
    else:
        print("Invalid choice. Please enter 1 or 2.")
        return get_input()

def read_file(filepath):
    if filepath.endswith('.txt'):
        with open(filepath, 'r') as file:
            return file.read()
    elif filepath.endswith('.docx'):
        return read_docx(filepath)
    elif filepath.endswith('.pdf'):
        return read_pdf(filepath)
    else:
        print("Unsupported file type. Please upload a TXT, DOCX, or PDF file.")
        return get_input()

def read_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_pdf(filepath):
    reader = PdfReader(filepath)
    full_text = []
    for page in reader.pages:
        full_text.append(page.extract_text())
    return '\n'.join(full_text)

def summarize_text(text, length):
    try:
        prompt = f"Summarize the following text in a {length} format: {text}"
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error in summarization: {e}")
        return None

def word_count(text):
    return len(text.split())

def save_as_pdf(text, filename):
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    content = []
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        content.append(Paragraph(paragraph, styles['Normal']))
    doc.build(content)

def save_as_docx(text, filename):
    doc = Document()
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)
    doc.save(filename)

def save_as_txt(text, filename):
    with open(filename, 'w') as file:
        file.write(text)

def process_summary(original_text, summary):
    original_word_count = word_count(original_text)
    summary_word_count = word_count(summary)

    print("\n" + "=" * 50)
    print("Summary:")
    print(summary)
    print("-" * 50)
    print(f"Original word count: {original_word_count}")
    print(f"Summary word count: {summary_word_count}")
    print(f"Reduction: {original_word_count - summary_word_count} words ({(1 - summary_word_count/original_word_count)*100:.2f}%)")
    print("=" * 50)

def handle_download_options(summary):
    download_choice = input("\nDo you want to download the summary? (Y/N): ").lower()
    if download_choice == 'y':
        format_choice = input("Choose format (PDF/DOCX/TXT): ").lower()
        if format_choice == 'pdf':
            filename = input("Enter filename for PDF: ")
            filename = filename if filename.endswith('.pdf') else filename + '.pdf'
            save_as_pdf(summary, filename)
            print(f"Summary saved as {filename}")
        elif format_choice == 'docx':
            filename = input("Enter filename for DOCX: ")
            filename = filename if filename.endswith('.docx') else filename + '.docx'
            save_as_docx(summary, filename)
            print(f"Summary saved as {filename}")
        elif format_choice == 'txt':
            filename = input("Enter filename for TXT: ")
            filename = filename if filename.endswith('.txt') else filename + '.txt'
            save_as_txt(summary, filename)
            print(f"Summary saved as {filename}")
        else:
            print("Invalid format choice. Please try again.")
    elif download_choice == 'n':
        print("Download skipped.")
    else:
        print("Invalid choice. Please enter Y or N.")

def main():
    while True:
        original_text = get_input()

        length_options = {'s': 'short', 'm': 'medium', 'l': 'long'}
        length = input("\nChoose summary length (S/M/L): ").lower()
        while length not in length_options:
            length = input("Invalid choice. Please enter S, M, or L: ").lower()

        summary = summarize_text(original_text, length_options[length])
        if summary:
            process_summary(original_text, summary)

            copy_choice = input("\nDo you want to copy the summary to clipboard? (Y/N): ").lower()
            if copy_choice == 'y':
                pyperclip.copy(summary)
                print("Summary copied to clipboard")

            handle_download_options(summary)

        continue_choice = input("\nDo you want to summarize another text? (Y/N): ").lower()
        if continue_choice != 'y':
            print("Thank you for using the summarizer. Goodbye!")
            break

if __name__ == "__main__":
    main()

# Tony Stark was a billionaire genius, known for his extravagant lifestyle and cutting-edge technology. As the head of Stark Industries, he revolutionized the world with his innovative weapons and gadgets. But his life took a drastic turn during a demonstration in Afghanistan. Captured by terrorists, Tony was mortally wounded by shrapnel near his heart. Imprisoned in a cave, he was coerced to build a missile for his captors. Instead, with the help of fellow prisoner Yinsen, he crafted a miniaturized arc reactor to power an electromagnet keeping the shrapnel from reaching his heart. They also secretly built a suit of powered armor, the first Iron Man suit. Escaping captivity, Tony returned to the United States a changed man. He announced Stark Industries would cease producing weapons, shocking the world and his business partner, Obadiah Stane. Tony focused on refining his suit, creating a sleek, powerful version equipped with advanced weaponry and flight capabilities. Adopting the mantle of Iron Man, Tony began fighting crime and terrorism, becoming a beacon of hope. He uncovered Stane's treachery; Stane had been selling weapons to terrorists and orchestrated Tony's capture. In a climactic battle, Iron Man defeated Stane, securing his place as a hero. Tony publicly revealed his identity as Iron Man, embracing his role as a protector. With his intelligence, courage, and heart, both literal and metaphorical, he embarked on a journey to make the world a safer place, forever leaving his mark as Iron Man.
# C:\Users\HAKIMI\Downloads\tortoise.txt